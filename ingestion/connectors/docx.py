"""Connecteur DOCX basé sur python-docx."""
from __future__ import annotations

from pathlib import Path
from typing import Iterable

from ingestion.config import ConnectorConfig
from ingestion.connectors.base import BaseConnector, DocumentChunk
from ingestion.metadata_utils import extract_ao_metadata, should_exclude_path


class DocxConnector(BaseConnector):
    """Découpe les documents DOCX par paragraphe."""

    document_type = "docx"

    def _iter_docx_files(self, directory: Path) -> Iterable[Path]:
        iterator = directory.rglob("*.docx") if self.config.recursive else directory.glob("*.docx")
        for candidate in iterator:
            if candidate.name.startswith("~$"):
                continue
            yield candidate

    def discover(self) -> Iterable[Path]:
        for path in self.config.paths:
            if path.is_dir():
                for candidate in self._iter_docx_files(path):
                    if should_exclude_path(candidate, self.config):
                        continue
                    yield candidate
            elif path.suffix.lower() == ".docx" and not path.name.startswith("~$"):
                if not should_exclude_path(path, self.config):
                    yield path

    def load(self, path: Path) -> Iterable[DocumentChunk]:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - dépendance optionnelle
            raise ImportError("python-docx doit être installé pour utiliser le connecteur DOCX") from exc

        document = Document(str(path))
        full_text = []

        # Paragraphes hors tableaux
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                full_text.append(text)

        # Contenu des tableaux converti en texte exploitable
        for table in document.tables:
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_rows.append(cells)

            if not table_rows:
                continue

            summary = self._summarize_table(table_rows)
            if summary:
                full_text.append(summary)
            else:
                for row in table_rows:
                    row_text = " | ".join(cell for cell in row if cell)
                    if row_text:
                        full_text.append(row_text)
        
        if not full_text:
            return

        joined_text = "\n".join(full_text)
        metadata = {
            "source": str(path),
            "document_type": self.document_type,
        }
        metadata.update(extract_ao_metadata(path))
        yield DocumentChunk(id=path.stem, text=joined_text, metadata=metadata)

    @staticmethod
    def _summarize_table(rows):
        """Essaye de produire une synthèse lisible pour les tableaux simples."""
        if not rows:
            return None

        header = rows[0]
        header_lower = [cell.lower() for cell in header]
        if len(header) < 2:
            return None

        first_cell = header_lower[0]
        # Cas particulier des tableaux financiers "En M€ 2022 2023 2024..."
        if "m€" in first_cell or "m€" in header[0]:
            years = header[1:]
            if not years:
                return None
            sentences = []
            unit = header[0].strip()
            for row in rows[1:]:
                label = row[0].strip()
                if not label:
                    continue
                parts = []
                for year, value in zip(years, row[1:]):
                    year = year.strip()
                    value = value.strip()
                    if not year or not value:
                        continue
                    parts.append(f"{value} {unit} en {year}")
                if parts:
                    sentences.append(f"{label} : {' ; '.join(parts)}.")
            if sentences:
                return " ".join(sentences)

        effectif_columns = [idx for idx, cell in enumerate(header_lower) if "effectif" in cell]
        if effectif_columns:
            label_idx = 0
            sentences = []
            for row in rows[1:]:
                if label_idx >= len(row):
                    continue
                label = row[label_idx].strip()
                if not label:
                    continue
                values = []
                for idx in effectif_columns:
                    if idx < len(row):
                        value = row[idx].strip()
                        if value:
                            values.append(value)
                if not values:
                    continue
                value_text = " / ".join(values)
                plural = "personnes" if value_text != "1" else "personne"
                sentences.append(f"{label} compte {value_text} {plural} (effectif).")
            if sentences:
                return " ".join(sentences)
        return None


__all__ = ["DocxConnector"]
